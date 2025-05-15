import streamlit as st
import requests
import os
import docx
import PyPDF2
from io import BytesIO
from docx import Document
from datetime import datetime
from dotenv import load_dotenv

# Load environment variables
load_dotenv()
GROQ_API_KEY = os.getenv('GROQ_API_KEY')
HF_TOKEN = os.getenv('HF_TOKEN')
JOOBLE_API_KEY = os.getenv('JOOBLE_API_KEY')

# Streamlit page setup
st.set_page_config(page_title="AI Career Assistant", layout="wide")
st.title("🚀 SmartHire : AI JOB PORTAL")

# Helper functions
def extract_docx(file):
    try:
        doc = docx.Document(file)
        return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    except Exception:
        return ""

def extract_pdf(file):
    try:
        reader = PyPDF2.PdfReader(file)
        return "".join([page.extract_text() or "" for page in reader.pages])
    except Exception:
        return ""

def query_groq(prompt):
    if len(prompt) > 5000:
        prompt = prompt[:5000]
    headers = {
        "Authorization": f"Bearer {GROQ_API_KEY}",
        "Content-Type": "application/json"
    }
    data = {
        "model": "llama3-70b-8192",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.5
    }
    response = requests.post("https://api.groq.com/openai/v1/chat/completions", headers=headers, json=data)
    if response.ok:
        return response.json()['choices'][0]['message']['content']
    else:
        return None

def get_jooble_jobs(keywords, experience_years, location):
    url = f"https://jooble.org/api/{JOOBLE_API_KEY}"
    payload = {'keywords': keywords, 'experience': experience_years, 'location': location}
    response = requests.post(url, json=payload)
    if response.ok:
        jobs = response.json().get('jobs', [])
        return [{
            'title': job['title'],
            'company': job['company'],
            'location': job['location'],
            'link': job.get('link', '')
        } for job in jobs]
    else:
        return []

def export_docx(text):
    doc = Document()
    for line in text.strip().split('\n'):
        doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# File upload and extraction
resume = st.file_uploader("Upload your resume (PDF, DOCX, or TXT)")
resume_content = ""

if resume:
    ext = resume.name.split('.')[-1].lower()
    if ext == 'pdf':
        resume_content = extract_pdf(resume)
    elif ext == 'docx':
        resume_content = extract_docx(resume)
    elif ext == 'txt':
        resume_content = resume.read().decode(errors='ignore')
    else:
        st.error("Unsupported file format. Please upload PDF, DOCX, or TXT.")

    if resume_content.strip():
        if len(resume_content) > 4000:
            st.warning("⚠ Resume is large; trimming to 4000 characters.")
            resume_content = resume_content[:4000]

        st.success("✅ Resume parsed successfully!")

        graduation_year = st.text_input("Graduation year (e.g., 2023)")
        stream = st.text_input("Graduation stream (e.g., Computer Science)")
        expected_salary = st.text_input("Expected salary (in USD or your currency)")
        location = st.text_input("Preferred job location (e.g., New York, Remote)")

        current_year = datetime.now().year
        try:
            grad_year_int = int(graduation_year)
            auto_experience = max(1, current_year - grad_year_int)
        except:
            auto_experience = 1

        # Suggested Jobs
        st.subheader("🔍 Suggested Jobs")
        if st.button("Get Recommended Jobs"):
            with st.spinner("Fetching recommendations..."):
                jooble_jobs = get_jooble_jobs(stream, auto_experience, location)
                if jooble_jobs:
                    for idx, job in enumerate(jooble_jobs, 1):
                        st.markdown(f"{idx}.** [{job['title']} at {job['company']} ({job['location']})]({job['link']})")
                else:
                    st.info("No suitable jobs found at this time.")

        # Resume Improvement Tips (NEW)
        st.subheader("✍ Resume Improvement Tips")
        job_position = st.text_input("Which position are you applying for?")
        if st.button("Get Resume Tips"):
            if job_position:
                tips_prompt = f"Give resume improvement suggestions for someone applying to a {job_position} position. Here's the resume:\n{resume_content}"
                with st.spinner("Generating tips..."):
                    tips = query_groq(tips_prompt)
                if tips:
                    st.write(tips)
                else:
                    st.error("Failed to generate tips. Please try again.")
            else:
                st.warning("Please specify the job position you're applying for.")

        # Cover Letter Generator (NEW)
        st.subheader("📝 Generate Cover Letter")
        company_name = st.text_input("What company are you applying to?")
        word_limit = st.slider("Desired word count for the cover letter", min_value=300, max_value=500, value=350, step=10)
        if st.button("Generate Cover Letter"):
            if company_name:
                cover_prompt = (
                    f"Generate a professional and tailored cover letter around {word_limit} words "
                    f"for the following resume. The user is applying to {company_name}:\n{resume_content}"
                )
                with st.spinner("Generating cover letter..."):
                    cover_letter = query_groq(cover_prompt)
                if cover_letter:
                    st.download_button("Download Cover Letter", export_docx(cover_letter), file_name="cover_letter.docx")
                else:
                    st.error("Failed to generate cover letter. Please try again.")
            else:
                st.warning("Please enter the company name you're applying to.")

        # Career Roadmap
        st.subheader("🗺 Career Roadmap")
        if st.button("Get Career Roadmap"):
            roadmap = query_groq(f"Generate a career roadmap for someone with this background:\n{resume_content}")
            if roadmap:
                st.write(roadmap)

        # Career Chatbot
        st.subheader("💬 Career Chatbot")
        user_question = st.text_input("Ask the AI Career Assistant a question:")
        if user_question:
            answer = query_groq(f"Question: {user_question}\nBased on this resume:\n{resume_content}")
            if answer:
                st.write(answer)

    else:
        st.error("Could not extract text from the uploaded resume.")
else:
    st.info("☝ Please upload your resume to get started.")
